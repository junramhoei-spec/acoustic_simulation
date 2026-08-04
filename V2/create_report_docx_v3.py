import os
import sys
import json
import shutil
import numpy as np
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

curr_dir = os.path.dirname(os.path.abspath(__file__))
res_v3_path = os.path.join(curr_dir, 'measured_plots_v3', 'results_v3.json')
plot_dir = os.path.join(curr_dir, 'measured_plots_v3')

with open(res_v3_path, encoding='utf-8') as f:
    data_v3 = json.load(f)

all_items = data_v3["all"]
groups_dict = data_v3["groups"]

# Also copy plots and results to artifact directory
art_dir = r'C:\Users\doilm\.gemini\antigravity-ide\brain\840a70b0-8f65-4630-9257-ba6424148a8a'
art_plot_dir = os.path.join(art_dir, 'measured_plots_v3')
os.makedirs(art_plot_dir, exist_ok=True)

for item in all_items:
    img_name = os.path.basename(item['img_path'])
    shutil.copy(item['img_path'], os.path.join(art_plot_dir, img_name))

doc = docx.Document()

# Page Margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Set base font
style_normal = doc.styles['Normal']
font = style_normal.font
font.name = 'Malgun Gothic'
font.size = Pt(10)
font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = title_p.add_run("🔊 실측 데이터셋 20종 전문가 결합(sline 앙상블)\n시간 순 그룹별(10개·3개·7개) 역추적 종합 평가 보고서 (v3)")
run_title.font.name = 'Malgun Gothic'
run_title.font.size = Pt(17)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# Calculate group-wise error metrics
g1_items = groups_dict["1"]["items"]
g2_items = groups_dict["2"]["items"]
g3_items = groups_dict["3"]["items"]

g1_h_err, g1_r_err = np.mean([i['err_h_mm'] for i in g1_items]), np.mean([i['err_r_mm'] for i in g1_items])
g2_h_err, g2_r_err = np.mean([i['err_h_mm'] for i in g2_items]), np.mean([i['err_r_mm'] for i in g2_items])
g3_h_err, g3_r_err = np.mean([i['err_h_mm'] for i in g3_items]), np.mean([i['err_r_mm'] for i in g3_items])

tot_h_err = np.mean([i['err_h_mm'] for i in all_items])
tot_r_err = np.mean([i['err_r_mm'] for i in all_items])

# Summary Box Table
summary_box = doc.add_table(rows=1, cols=1)
summary_box.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = summary_box.cell(0, 0)
cell.width = Inches(6.5)

tcPr = cell._element.get_or_add_tcPr()
shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F9FF"/>')
tcPr.append(shd)

borders = parse_xml(
    f'<w:tcBorders {nsdecls("w")}>\n'
    f'  <w:top w:val="none"/>\n'
    f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="1D4ED8"/>\n'
    f'  <w:bottom w:val="none"/>\n'
    f'  <w:right w:val="none"/>\n'
    f'</w:tcBorders>'
)
tcPr.append(borders)

p_sum = cell.paragraphs[0]
p_sum.paragraph_format.space_before = Pt(6)
p_sum.paragraph_format.space_after = Pt(6)
p_sum.paragraph_format.left_indent = Inches(0.1)

r_sum = p_sum.add_run("📌 시간 순 그룹별 핵심 평가지표 요약\n")
r_sum.font.bold = True
r_sum.font.size = Pt(11)
r_sum.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

metrics_text = (
    f"• [1그룹] 7월 실측 세션 (10개): H MAE = {g1_h_err:.2f} mm | r MAE = {g1_r_err:.2f} mm\n"
    f"• [2그룹] 8월 4일 오전 사전 세션 (3개): H MAE = {g2_h_err:.2f} mm | r MAE = {g2_r_err:.2f} mm\n"
    f"• [3그룹] 8월 4일 오후 정밀 링 컵 세션 (7개): H MAE = {g3_h_err:.2f} mm | r MAE = {g3_r_err:.2f} mm\n"
    f"----------------------------------------------------------------------\n"
    f"• [전체 20개 총합]: H MAE = {tot_h_err:.2f} mm | r MAE = {tot_r_err:.2f} mm\n"
    f"• 적용 모델: H 전문가 (rnn_sline_nodetach_bestH.pt) + r 전문가 (rnn_sline_uni.pt)"
)
r_m = p_sum.add_run(metrics_text)
r_m.font.size = Pt(9.5)

doc.add_paragraph().paragraph_format.space_after = Pt(10)

def add_group_table_and_grid(g_key, g_title):
    g_data = groups_dict[g_key]
    items = g_data["items"]

    h1 = doc.add_heading(f"📊 {g_title}", level=1)
    h1.runs[0].font.name = 'Malgun Gothic'
    h1.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    # Group Metrics Table
    table = doc.add_table(rows=len(items) + 1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["No", "실측 대상 컵 및 도넛 링 조합", "참값 H", "예측 H", "H 오차", "r 오차", "예측 부피", "스텝"]
    col_widths = [Inches(0.4), Inches(2.3), Inches(0.6), Inches(0.6), Inches(0.6), Inches(0.6), Inches(0.7), Inches(0.5)]

    hdr_cells = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr_cells[idx].width = col_widths[idx]
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tcPr = hdr_cells[idx]._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E3A8A"/>')
        tcPr.append(shd)

    for idx, r_data in enumerate(items, 1):
        row_cells = table.rows[idx].cells
        row_values = [
            f"#{r_data['global_idx']:02d}",
            r_data['name'],
            f"{r_data['gt_h_mm']:.1f}mm",
            f"{r_data['pred_h_mm']:.1f}mm",
            f"{r_data['err_h_mm']:.1f}mm",
            f"{r_data['err_r_mm']:.1f}mm",
            f"{r_data['vol_pred_ml']:.0f}mL",
            f"{r_data['steps']}s"
        ]
        bg_color = "F9FAFB" if idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_values):
            row_cells[c_idx].width = col_widths[c_idx]
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(val)
            run.font.size = Pt(8.5)
            tcPr = row_cells[c_idx]._element.get_or_add_tcPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
            tcPr.append(shd)

    for row in table.rows:
        for cell_item in row.cells:
            tcPr = cell_item._element.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>\n'
                f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>\n'
                f'  <w:left w:val="none"/>\n'
                f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>\n'
                f'  <w:right w:val="none"/>\n'
                f'</w:tcBorders>'
            )
            tcPr.append(borders)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 2-Column Grid Images for this group
    grid_table = doc.add_table(rows=(len(items) + 1) // 2, cols=2)
    grid_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i in range(0, len(items), 2):
        row_idx = i // 2
        row_cells = grid_table.rows[row_idx].cells

        # Left cell
        r_left = items[i]
        p_l = row_cells[0].paragraphs[0]
        p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_l.paragraph_format.space_after = Pt(2)
        p_l.add_run().add_picture(r_left['img_path'], width=Inches(3.0))
        p_lt = row_cells[0].add_paragraph()
        p_lt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_lt.paragraph_format.space_after = Pt(8)
        r_lt_txt = p_lt.add_run(f"#{r_left['global_idx']:02d}. {r_left['name']}\n(H 오차 {r_left['err_h_mm']:.1f}mm / r 오차 {r_left['err_r_mm']:.1f}mm)")
        r_lt_txt.font.size = Pt(8.0)
        r_lt_txt.font.bold = True

        # Right cell
        if i + 1 < len(items):
            r_right = items[i + 1]
            p_r = row_cells[1].paragraphs[0]
            p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_r.paragraph_format.space_after = Pt(2)
            p_r.add_run().add_picture(r_right['img_path'], width=Inches(3.0))
            p_rt = row_cells[1].add_paragraph()
            p_rt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_rt.paragraph_format.space_after = Pt(8)
            r_rt_txt = p_rt.add_run(f"#{r_right['global_idx']:02d}. {r_right['name']}\n(H 오차 {r_right['err_h_mm']:.1f}mm / r 오차 {r_right['err_r_mm']:.1f}mm)")
            r_rt_txt.font.size = Pt(8.0)
            r_rt_txt.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Build Sections for 3 groups
add_group_table_and_grid("1", "1그룹: 7월 실측 세션 (10개 — 유리/고무/3D프린터/텀블러/실린더)")
add_group_table_and_grid("2", "2그룹: 8월 4일 오전 사전 테스트 세션 (3개 — 150mm 일자컵/벌어진컵)")
add_group_table_and_grid("3", "3그룹: 8월 4일 오후 정밀 링 컵 실측 세션 (7개 — 15mm 링 계단)")

# Conclusions
h_c = doc.add_heading("💡 시간 순 그룹별 분석 결론 및 시사점", level=1)
h_c.runs[0].font.name = 'Malgun Gothic'
h_c.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

p_c = doc.add_paragraph()
p_c.add_run("1. 7월 실측 세션 (1그룹, 10개): H MAE = 17.06 mm / r MAE = 6.11 mm\n").font.bold = True
p_c.add_run("유리, 고무, 3D프린터 등 이종 재질 컵들이 포함되어 있으며, 재질 및 컵 형상 차이에도 불구하고 반지름 오차가 6.1mm 수준으로 우수한 역추적 정확도를 유지합니다.\n\n")

p_c.add_run("2. 8월 4일 오전 사전 세션 (2그룹, 3개): H MAE = 5.71 mm / r MAE = 4.42 mm\n").font.bold = True
p_c.add_run("150mm 대형 일자 컵 및 초초검검파파 벌어진 컵에 대한 40~50mL 대단위 주입 세션으로, 높이 오차가 5.7mm로 매우 정밀하게 복원되었습니다.\n\n")

p_c.add_run("3. 8월 4일 오후 정밀 링 컵 세션 (3그룹, 7개): H MAE = 15.09 mm / r MAE = 5.50 mm\n").font.bold = True
p_c.add_run("15mm 두께의 색상 도넛 링으로 구성된 다단계 계단/볼록/오므라드는 complex 프로파일에 대한 실측 세션으로, 계단 변화에도 불구하고 반지름 오차 5.5mm 내에서 전체 형상 윤곽을 정교하게 추적했습니다.")

docx_v3_path = os.path.join(curr_dir, '실측_20종_시간순그룹별_역추적_종합평가보고서_v3.docx')
alt_v3_path = os.path.join(curr_dir, '실측_20종_시간순그룹별_역추적_종합평가보고서_v3_out.docx')

try:
    doc.save(docx_v3_path)
    print(f"Successfully generated DOCX v3 at: {docx_v3_path}")
    saved_path = docx_v3_path
except PermissionError:
    doc.save(alt_v3_path)
    print(f"File was locked. Saved to alternative path: {alt_v3_path}")
    saved_path = alt_v3_path

try:
    shutil.copy(saved_path, os.path.join(art_dir, '실측_20종_시간순그룹별_역추적_종합평가보고서_v3.docx'))
    print("Copied DOCX v3 to artifact folder!")
except Exception as e:
    print(f"Artifact copy skipped: {e}")


