import os
import sys
import json
import numpy as np
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

curr_dir = os.path.dirname(os.path.abspath(__file__))

# Path to results.json and plots
res_path = os.path.join(curr_dir, 'measured_plots', 'results.json')
plot_dir = os.path.join(curr_dir, 'measured_plots')

with open(res_path, encoding='utf-8') as f:
    results = json.load(f)

doc = docx.Document()

# Page Margins (1 inch = 72 pt)
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Set base font
style_normal = doc.styles['Normal']
font = style_normal.font
font.name = 'Malgun Gothic'
font.size = Pt(10)
font.color.rgb = RGBColor(0x1F, 0x29, 0x37) # Dark grey

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = title_p.add_run("🔊 실측 데이터셋 20종 전문가 결합(sline 앙상블)\n역추적 종합 평가 보고서")
run_title.font.name = 'Malgun Gothic'
run_title.font.size = Pt(18)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Deep Navy

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# Summary Box Table
summary_box = doc.add_table(rows=1, cols=1)
summary_box.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = summary_box.cell(0, 0)
cell.width = Inches(6.5)

# Set cell background to light blue (#F0F9FF) and left border dark blue
tcPr = cell._element.get_or_add_tcPr()
shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F9FF"/>')
tcPr.append(shd)

borders = parse_xml(
    f'<w:tcBorders {nsdecls("w")}>\n'
    f'  <w:top w:val="none"/>\n'
    f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="1D4ED8"/>\n'
    f'  <w:bottom w:val="none"/>\n'
    f'  <w:right w:val="none"/>\n'
    f'</w:tcBorders>'
)
tcPr.append(borders)

p_sum = cell.paragraphs[0]
p_sum.paragraph_format.space_before = Pt(6)
p_sum.paragraph_format.space_after = Pt(6)
p_sum.paragraph_format.left_indent = Inches(0.1)

r_sum = p_sum.add_run("📌 핵심 평가지표 요약\n")
r_sum.font.bold = True
r_sum.font.size = Pt(11)
r_sum.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

h_errs = [r['err_h_mm'] for r in results]
r_errs = [r['err_r_mm'] for r in results]

metrics_text = (
    f"• 평가 세션 수: 총 20개 실측 컵 데이터 (총 309 스텝 물채움)\n"
    f"• 평균 높이 오차 (H MAE): {np.mean(h_errs):.2f} mm (중앙값: {np.median(h_errs):.2f} mm)\n"
    f"• 평균 반지름 오차 (r MAE): {np.mean(r_errs):.2f} mm (중앙값: {np.median(r_errs):.2f} mm)\n"
    f"• 적용 모델: H 전문가 (rnn_sline_nodetach_bestH.pt) + r 전문가 (rnn_sline_uni.pt)"
)
r_m = p_sum.add_run(metrics_text)
r_m.font.size = Pt(10)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Heading 1: Summary Table
h1 = doc.add_heading("📊 1. 20개 실측 데이터 세션 종합 요약표", level=1)
h1.runs[0].font.name = 'Malgun Gothic'
h1.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

table = doc.add_table(rows=len(results) + 1, cols=8)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

headers = ["No", "실측 대상 컵 및 도넛 링 조합", "참값 H", "예측 H", "H 오차", "r 오차", "예측 부피", "스텝"]
col_widths = [Inches(0.4), Inches(2.3), Inches(0.6), Inches(0.6), Inches(0.6), Inches(0.6), Inches(0.7), Inches(0.5)]

# Format Header Row
hdr_cells = table.rows[0].cells
for idx, text in enumerate(headers):
    hdr_cells[idx].width = col_widths[idx]
    p = hdr_cells[idx].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    tcPr = hdr_cells[idx]._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E3A8A"/>')
    tcPr.append(shd)

# Data Rows
for idx, r_data in enumerate(results, 1):
    row_cells = table.rows[idx].cells
    row_values = [
        f"{r_data['idx']:02d}",
        r_data['name'],
        f"{r_data['gt_h_mm']:.1f}mm",
        f"{r_data['pred_h_mm']:.1f}mm",
        f"{r_data['err_h_mm']:.1f}mm",
        f"{r_data['err_r_mm']:.1f}mm",
        f"{r_data['vol_pred_ml']:.0f}mL",
        f"{r_data['steps']}s"
    ]
    bg_color = "F9FAFB" if idx % 2 == 1 else "FFFFFF"
    for c_idx, val in enumerate(row_values):
        row_cells[c_idx].width = col_widths[c_idx]
        p = row_cells[c_idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(val)
        run.font.size = Pt(8.5)
        tcPr = row_cells[c_idx]._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
        tcPr.append(shd)

# Set borders for table
for row in table.rows:
    for cell in row.cells:
        tcPr = cell._element.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>\n'
            f'  <w:left w:val="none"/>\n'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>\n'
            f'  <w:right w:val="none"/>\n'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Heading 2: 2-Column Grid Images
h2 = doc.add_heading("🖼️ 2. 2열 그리드 형상 비교 프로파일 그래프 (True vs Pred)", level=1)
h2.runs[0].font.name = 'Malgun Gothic'
h2.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

p_note = doc.add_paragraph()
r_note = p_note.add_run("파란색 라인/채우기: 참값 원본 형상 (True Ground Truth)\n빨간색 라인/채우기: sline 앙상블 모델 예측 형상 (Pred Ensemble)")
r_note.font.size = Pt(9)
r_note.font.italic = True
r_note.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

grid_table = doc.add_table(rows=(len(results) + 1) // 2, cols=2)
grid_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i in range(0, len(results), 2):
    row_idx = i // 2
    row_cells = grid_table.rows[row_idx].cells

    # Left cell
    r_left = results[i]
    p_l = row_cells[0].paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_l.paragraph_format.space_after = Pt(2)
    p_l.add_run().add_picture(r_left['img_path'], width=Inches(3.1))
    p_lt = row_cells[0].add_paragraph()
    p_lt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_lt.paragraph_format.space_after = Pt(12)
    r_lt_txt = p_lt.add_run(f"#{r_left['idx']:02d}. {r_left['name']}\n(H 오차 {r_left['err_h_mm']:.1f}mm / r 오차 {r_left['err_r_mm']:.1f}mm)")
    r_lt_txt.font.size = Pt(8.5)
    r_lt_txt.font.bold = True

    # Right cell
    if i + 1 < len(results):
        r_right = results[i + 1]
        p_r = row_cells[1].paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_r.paragraph_format.space_after = Pt(2)
        p_r.add_run().add_picture(r_right['img_path'], width=Inches(3.1))
        p_rt = row_cells[1].add_paragraph()
        p_rt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_rt.paragraph_format.space_after = Pt(12)
        r_rt_txt = p_rt.add_run(f"#{r_right['idx']:02d}. {r_right['name']}\n(H 오차 {r_right['err_h_mm']:.1f}mm / r 오차 {r_right['err_r_mm']:.1f}mm)")
        r_rt_txt.font.size = Pt(8.5)
        r_rt_txt.font.bold = True

# Heading 3: Conclusions
doc.add_page_break()
h3 = doc.add_heading("💡 3. 역추적 분석 결론 및 시사점", level=1)
h3.runs[0].font.name = 'Malgun Gothic'
h3.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

p_c1 = doc.add_paragraph()
p_c1.add_run("1. 반지름 오차의 높은 안정성 (r MAE ≈ 5.64 mm)\n").font.bold = True
p_c1.add_run("다양한 재질(유리, 3D프린터 ABS, 고무, 텀블러)과 형상(일자형, 벌어진형, 오므라든형, 볼록형)에 관계없이 반지름 역추적이 평균 5.6mm 이내로 매우 정밀하게 재구성되었습니다.\n\n")

p_c2 = doc.add_paragraph()
p_c2.add_run("2. 높이 추정 정확도 (H MAE ≈ 14.67 mm)\n").font.bold = True
p_c2.add_run("유리 일자 컵(#17)의 경우 오차 0.3 mm, 일자 컵 150mm(#18, #19)의 경우 오차 4.5 mm ~ 4.6 mm로 뛰어난 정확도를 보였습니다. 컵 상단부 공기기둥 남김 비율이나 주입 스텝 간격에 따른 편차가 일부 존재하나, 전체적으로 컵의 개략적 높이와 형상 경향성을 우수하게 복원합니다.")

docx_out_path = os.path.join(curr_dir, '실측_20종_전문가결합_역추적_종합평가보고서.docx')
try:
    doc.save(docx_out_path)
    print(f"Successfully generated docx report at: {docx_out_path}")
except PermissionError:
    alt_path = os.path.join(curr_dir, '실측_20종_전문가결합_역추적_종합평가보고서_v2.docx')
    doc.save(alt_path)
    print(f"File was locked by Word. Saved to alternative path: {alt_path}")
